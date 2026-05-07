"""Generate a pre-filled Joining Kit Word (.docx) document for a selected candidate.
Mirrors the 14-section structure of joining_kit_pdf.py. Word natively handles
Devanagari complex-script shaping (GSUB/GPOS), so Nirmala UI renders correctly.
"""
from io import BytesIO
from datetime import datetime
import os

from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Helpers — data extraction (same logic as joining_kit_pdf.py)
# ---------------------------------------------------------------------------

def _full_name(c: dict) -> str:
    return " ".join(filter(None, [c.get("first_name"), c.get("last_name")])).strip()


def _fmt_date(val: str | None) -> str:
    if not val:
        return "___/___/______"
    try:
        return datetime.strptime(val[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        return val or "___/___/______"


def _addr_line(c: dict) -> str:
    parts = filter(None, [
        c.get("address_line1"), c.get("address_line2"),
        c.get("city"), c.get("state"), c.get("pincode"),
    ])
    return ", ".join(parts)


def _blank(n: int = 20) -> str:
    return "_" * n


# ---------------------------------------------------------------------------
# Docx formatting helpers
# ---------------------------------------------------------------------------

_ENG_FONT  = "Times New Roman"
_HINDI_FONT = "Nirmala UI"
_BOLD_FONT  = "Times New Roman"
_HEADING_SZ = 13   # pt
_BODY_SZ    = 9    # pt
_SMALL_SZ   = 8    # pt


def _set_font(run, font_name: str, size_pt: int, bold: bool = False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # For Nirmala UI, also set the complex-script font
    rPr = run._r.get_or_add_rPr()
    cs = OxmlElement("w:rFonts")
    cs.set(qn("w:cs"), font_name)
    rPr.append(cs)


def _heading(doc: Document, text: str, level: int = 1, center: bool = True, hindi: bool = False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    font_name = _HINDI_FONT if hindi else _ENG_FONT
    sz = _HEADING_SZ if level == 1 else 11
    _set_font(run, font_name, sz, bold=not hindi)
    return p


def _para(doc: Document, text: str = "", bold: bool = False, center: bool = False,
          hindi: bool = False, size: int = _BODY_SZ, italic: bool = False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    font_name = _HINDI_FONT if hindi else _ENG_FONT
    _set_font(run, font_name, size, bold=bold)
    run.italic = italic
    return p


def _bilingual_para(doc: Document, en_text: str, hi_text: str, bullet: bool = False):
    """One paragraph: English text followed by Hindi text on a new line."""
    prefix = "• " if bullet else ""
    p = doc.add_paragraph()
    r1 = p.add_run(prefix + en_text + "\n")
    _set_font(r1, _ENG_FONT, _BODY_SZ)
    r2 = p.add_run(hi_text)
    _set_font(r2, _HINDI_FONT, _BODY_SZ)
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "40")
    pPr.append(spacing)
    return p


def _kv_table(doc: Document, rows: list[tuple[str, str]], col1_w: float = 65, col2_w: float = 105,
              hindi_key: bool = False):
    """Two-column key:value table. col widths in mm."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].width = Mm(col1_w)
        cells[1].width = Mm(col2_w)
        # Key cell
        kp = cells[0].paragraphs[0]
        kr = kp.add_run(key)
        _set_font(kr, _HINDI_FONT if hindi_key else _ENG_FONT, _SMALL_SZ, bold=not hindi_key)
        # Value cell
        vp = cells[1].paragraphs[0]
        vr = vp.add_run(str(val))
        _set_font(vr, _ENG_FONT, _SMALL_SZ)
    return table


def _sig_table(doc: Document, labels: list[str]):
    """Horizontal signature block with equal columns."""
    n = len(labels)
    table = doc.add_table(rows=2, cols=n)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = Mm(170 / n)
    for i, label in enumerate(labels):
        for r in range(2):
            table.rows[r].cells[i].width = col_w
        lp = table.rows[0].cells[i].paragraphs[0]
        lr = lp.add_run(label)
        _set_font(lr, _ENG_FONT, _SMALL_SZ, bold=True)
        vp = table.rows[1].cells[i].paragraphs[0]
        vr = vp.add_run("")
        _set_font(vr, _ENG_FONT, _SMALL_SZ)
    return table


def _grid_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Multi-column table with bold header row."""
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = Mm(170 / n)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].width = col_w
        hp = table.rows[0].cells[i].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(h)
        _set_font(hr, _ENG_FONT, _SMALL_SZ, bold=True)
    for ri, row in enumerate(rows):
        for ci, cell_val in enumerate(row):
            table.rows[ri + 1].cells[ci].width = col_w
            cp = table.rows[ri + 1].cells[ci].paragraphs[0]
            cr = cp.add_run(cell_val)
            _set_font(cr, _ENG_FONT, _SMALL_SZ)
    return table


def _spacer(doc: Document, lines: int = 1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def _page_break(doc: Document):
    doc.add_page_break()


def _section_divider(doc: Document, title: str):
    """Centered horizontal rule + section number. Used between section headings."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    _set_font(run, _ENG_FONT, 7)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _s1_checklist(doc: Document, c: dict, company: dict,
                  has_aadhaar: bool = False, has_pan: bool = False):
    co = company.get("name", "Radhya Micro Finance Private Limited")
    co_addr = company.get("address", "Head Office, Moradabad (U.P.)")
    _heading(doc, co.upper(), level=1)
    _para(doc, co_addr, center=True, size=_SMALL_SZ, italic=True)
    _spacer(doc)
    _heading(doc, "JOINING KIT — DOCUMENT CHECKLIST", level=1)
    _spacer(doc)
    items = [
        ("1",  "Employee Information Sheet",                    True,  "Included"),
        ("2",  "Staff Undertaking",                             True,  "Included"),
        ("3",  "Employee Insurance Form",                       True,  "Included"),
        ("4",  "Gratuity Nomination Form (Form F)",             True,  "Included"),
        ("5",  "PF Declaration",                                True,  "Included"),
        ("6",  "Form 11 (PF)",                                  True,  "Included"),
        ("7",  "ESI Declaration",                               True,  "Included"),
        ("8",  "Notice Period Policy",                          True,  "Included"),
        ("9",  "Online Police Verification Report",             False, ""),
        ("10", "PF Proof Document",                             False, ""),
        ("11", "ESIC Proof Document",                           False, ""),
        ("12", "PAN Card Copy",                                 has_pan, "Already provided" if has_pan else ""),
        ("13", "Consent Form — Original Documents Submission",  True,  "Included"),
        ("14", "POSH Declaration & BGV Consent Form",           True,  "Included"),
    ]
    table = doc.add_table(rows=1 + len(items), cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Sr", "Document", "Status", "Remarks"]):
        hp = table.rows[0].cells[i].paragraphs[0]
        hr_ = hp.add_run(h)
        _set_font(hr_, _ENG_FONT, _SMALL_SZ, bold=True)
    for ri, (sr, doc_name, done, rem) in enumerate(items):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate([sr, doc_name, "✓ Done" if done else "Pending", rem]):
            cp = cells[ci].paragraphs[0]
            cr = cp.add_run(val)
            _set_font(cr, _ENG_FONT, _SMALL_SZ)


def _s2_employee_info(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "EMPLOYEE INFORMATION SHEET")
    _spacer(doc)
    name = _full_name(c) or _blank()
    dob  = _fmt_date(c.get("date_of_birth"))
    doj  = _fmt_date(c.get("expected_joining_date"))
    rows = [
        ("Employee ID",          c.get("employee_id") or _blank()),
        ("Full Name",            name),
        ("Date of Birth",        dob),
        ("Gender",               c.get("gender") or _blank()),
        ("Father / Husband Name", c.get("father_or_husband_name") or _blank()),
        ("Mobile",               c.get("phone") or _blank()),
        ("Email",                c.get("email") or _blank()),
        ("Aadhaar No.",          c.get("aadhaar_number") or _blank()),
        ("PAN No.",              c.get("pan_number") or _blank()),
        ("Address",              _addr_line(c) or _blank()),
        ("Designation",          c.get("position") or _blank()),
        ("Department",           c.get("department") or _blank()),
        ("Date of Joining",      doj),
        ("Bank Name",            c.get("bank_name") or _blank()),
        ("Account No.",          c.get("account_number") or _blank()),
        ("IFSC Code",            c.get("ifsc_code") or _blank()),
        ("UAN No.",              c.get("uan_number") or _blank()),
        ("ESI No.",              c.get("esi_number") or _blank()),
    ]
    _kv_table(doc, rows)
    _spacer(doc)

    # Nomination
    nom_hdr = [("Sr", "Name", "Relationship", "Date of Birth", "Share %")]
    _heading(doc, "Nominee Details", level=2, center=False)
    nominees = c.get("nominees") or [{}]
    nom_rows = []
    for i, n in enumerate(nominees[:4], 1):
        nom_rows.append([
            str(i),
            f"{n.get('first_name','')} {n.get('last_name','')}".strip() or _blank(15),
            n.get("relationship") or _blank(12),
            _fmt_date(n.get("date_of_birth")),
            n.get("share_percentage") or _blank(5),
        ])
    if not nom_rows:
        nom_rows = [["1", _blank(15), _blank(12), _blank(10), _blank(5)]]
    _grid_table(doc, ["Sr", "Name", "Relationship", "Date of Birth", "Share %"], nom_rows)
    _spacer(doc)

    # Educational
    _heading(doc, "Educational Qualifications", level=2, center=False)
    edu_rows = c.get("education") or []
    edu_tbl_rows = [[
        e.get("degree") or _blank(12),
        e.get("institution") or _blank(15),
        e.get("year_of_passing") or _blank(6),
        e.get("percentage") or _blank(5),
    ] for e in edu_rows[:6]]
    if not edu_tbl_rows:
        edu_tbl_rows = [[_blank(12), _blank(15), _blank(6), _blank(5)]]
    _grid_table(doc, ["Degree", "Institution", "Year", "%"], edu_tbl_rows)


def _s3_undertaking(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "Staff Undertaking")
    name  = _full_name(c) or _blank()
    fhn   = c.get("father_or_husband_name") or _blank()
    addr  = _addr_line(c) or _blank()
    doj   = _fmt_date(c.get("expected_joining_date"))
    _bilingual_para(doc,
        f"I, {name}, S/o or D/o or W/o Mr. {fhn}, residing at {addr}",
        "मैं पुत्र / पुत्री / पत्नी श्री ____ उम्र ____ वर्ष गाँव और पोस्ट मण्डल जिला ____ का निवासी हूँ।")
    _bilingual_para(doc,
        "Do hereby affirm on oath that:",
        "एतद् द्वारा शपथपूर्वक इसकी पुष्टि करता हूँ:")
    bullets = [
        ("The information regarding my academic qualifications and work experiences mentioned in my bio-data is true and accurate.",
         "मेरे बायोडाटा में दी गई सभी जानकारी सत्य है।"),
        ("There have never been any legal proceedings initiated against me, nor have I been involved in any misconduct / fraud / embezzlement of cash.",
         "मेरे खिलाफ कभी भी कोई कानूनी कार्यवाही नहीं हुई है और न ही मैं धोखाधड़ी / गबन आदि में शामिल रहा हूँ।"),
        ("I declare the above statements to be true in all respects. Any false statement may make me liable to dismissal without prior notice.",
         "मैं उपरोक्त सभी कथनों को सत्य घोषित करता हूँ। कोई भी गलत बयान मुझे बिना किसी पूर्व सूचना के बर्खास्त करवा सकता है।"),
    ]
    for en, hi in bullets:
        _bilingual_para(doc, en, hi, bullet=True)
    _bilingual_para(doc, "I further declare that —", "मैं यह और भी घोषित करता हूँ —")
    further = [
        ("Will carry out responsibilities with care, diligence and thoroughness per Job Description and Operations Manual.",
         "अपने कार्य विवरण और संचालन नियमावली के अनुसार अपनी ज़िम्मेदारियों को सावधानी से निभाऊँगा।"),
        ("Will be honest and sincere in all work and set a good example to build the Company's public image.",
         "मैं अपने सभी कार्यों में ईमानदार रहूँगा और कंपनी की सार्वजनिक छवि बनाऊँगा।"),
        ("Will not engage in any business, service or occupation of any kind whatsoever during the period of employment.",
         "मैं रोजगार की अवधि के दौरान किसी भी प्रकार का व्यवसाय या अन्य नौकरी नहीं करूँगा।"),
        ("Will maintain the confidentiality of all information and materials related to the Company.",
         "मैं कंपनी से संबंधित सभी जानकारियों की गोपनीयता बनाए रखूँगा।"),
    ]
    for en, hi in further:
        _bilingual_para(doc, en, hi, bullet=True)
    _spacer(doc)
    _kv_table(doc, [
        ("Joining Location:", c.get("joining_location") or "Head Office, Moradabad"),
        ("Date of Joining:", doj),
        ("Employee Signature:", ""),
    ])


def _s4_insurance(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "EMPLOYEE'S INSURANCE FORM")
    _para(doc, "(कर्मचारी के बीमा पत्र)", center=True, hindi=True, size=10)
    _para(doc, "(Group Medical Insurance, Group Personal Accident & Group Term Life)", center=True, size=_SMALL_SZ, italic=True)
    _spacer(doc)
    name = _full_name(c) or _blank()
    doj  = _fmt_date(c.get("expected_joining_date"))
    _para(doc, f"I hereby undertake that I wish to become a member of the Employee's Insurance Policy. "
               f"I, {name}, acknowledge that the premium for the policy is to be paid by the Company on my behalf.")
    _spacer(doc)
    noms = c.get("nominees") or [{}]
    nom_rows = []
    for i, n in enumerate(noms[:4], 1):
        nom_rows.append([
            str(i),
            f"{n.get('first_name','')} {n.get('last_name','')}".strip() or _blank(15),
            n.get("relationship") or _blank(10),
            _fmt_date(n.get("date_of_birth")),
            n.get("share_percentage") or _blank(4),
            n.get("address") or _blank(15),
        ])
    if not nom_rows:
        nom_rows = [["1", _blank(15), _blank(10), _blank(10), _blank(4), _blank(15)],
                    ["2", _blank(15), _blank(10), _blank(10), _blank(4), _blank(15)]]
    _grid_table(doc, ["Sr", "Nominee Name", "Relationship", "DOB", "Share%", "Address"], nom_rows)
    _spacer(doc)
    _kv_table(doc, [
        ("Name:", name),
        ("Date of Joining:", doj),
        ("Designation:", c.get("position") or _blank()),
        ("Signature:", ""),
        ("Date:", ""),
    ])


def _s5_gratuity(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "FORM F — GRATUITY NOMINATION")
    _para(doc, "(Under Section 6 of the Payment of Gratuity Act 1972)", center=True, size=_SMALL_SZ, italic=True)
    _spacer(doc)
    name = _full_name(c) or _blank()
    doj  = _fmt_date(c.get("expected_joining_date"))
    _para(doc, f"I, {name}, hereby nominate the following person(s) for receipt of gratuity payable after my death.")
    _spacer(doc)
    noms = c.get("nominees") or [{}]
    nom_rows = []
    for i, n in enumerate(noms[:4], 1):
        nom_rows.append([
            str(i),
            f"{n.get('first_name','')} {n.get('last_name','')}".strip() or _blank(15),
            n.get("relationship") or _blank(10),
            n.get("address") or _blank(20),
            _fmt_date(n.get("date_of_birth")),
            n.get("share_percentage") or _blank(4),
        ])
    if not nom_rows:
        nom_rows = [["1", _blank(15), _blank(10), _blank(20), _blank(10), _blank(4)],
                    ["2", _blank(15), _blank(10), _blank(20), _blank(10), _blank(4)]]
    _grid_table(doc, ["Sr", "Nominee Name", "Relation", "Address", "DOB", "Share%"], nom_rows)
    _spacer(doc)
    _kv_table(doc, [
        ("Name:", name),
        ("Designation:", c.get("position") or _blank()),
        ("Date of Joining:", doj),
        ("Address:", _addr_line(c) or _blank()),
        ("Signature of Employee:", ""),
        ("Date:", ""),
        ("Signature of Witness 1:", ""),
        ("Signature of Witness 2:", ""),
    ])


def _s6_pf(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "PROVIDENT FUND — DECLARATION")
    _spacer(doc)
    name = _full_name(c) or _blank()
    dob  = _fmt_date(c.get("date_of_birth"))
    _para(doc, f"I, {name}, hereby declare my details for Provident Fund registration:")
    _spacer(doc)
    _kv_table(doc, [
        ("Name:",                       name),
        ("Father/Husband Name:",        c.get("father_or_husband_name") or _blank()),
        ("Date of Birth:",              dob),
        ("Gender:",                     c.get("gender") or _blank()),
        ("Aadhaar No.:",                c.get("aadhaar_number") or _blank()),
        ("PAN No.:",                    c.get("pan_number") or _blank()),
        ("Mobile No.:",                 c.get("phone") or _blank()),
        ("Email:",                      c.get("email") or _blank()),
        ("Bank Name:",                  c.get("bank_name") or _blank()),
        ("Account No.:",                c.get("account_number") or _blank()),
        ("IFSC Code:",                  c.get("ifsc_code") or _blank()),
        ("UAN No. (if existing):",      c.get("uan_number") or _blank()),
        ("Previous PF Account (if any):", _blank()),
        ("International Worker:",       "Yes / No"),
        ("Employee Signature:",         ""),
        ("Date:",                       ""),
    ])


def _s7_form11(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "FORM 11 — EPF / EPS REVISED DECLARATION")
    _spacer(doc)
    name = _full_name(c) or _blank()
    dob  = _fmt_date(c.get("date_of_birth"))
    doj  = _fmt_date(c.get("expected_joining_date"))
    _kv_table(doc, [
        ("Name of Member:",             name),
        ("Father/Husband Name:",        c.get("father_or_husband_name") or _blank()),
        ("Date of Birth:",              dob),
        ("Gender:",                     c.get("gender") or _blank()),
        ("Relation:",                   "Father / Husband"),
        ("Mobile:",                     c.get("phone") or _blank()),
        ("Email:",                      c.get("email") or _blank()),
        ("Aadhaar No.:",                c.get("aadhaar_number") or _blank()),
        ("PAN No.:",                    c.get("pan_number") or _blank()),
        ("Date of Joining:",            doj),
        ("UAN (if existing):",          c.get("uan_number") or _blank()),
        ("Previous PF A/C No.:",        _blank()),
        ("Pension Scheme Member:",      "Yes / No"),
        ("KYC Verified:",               "Yes / No"),
        ("Employee Signature:",         ""),
        ("Date:",                       ""),
    ])


def _s8_esi(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "ESIC DECLARATION FORM")
    _spacer(doc)
    name = _full_name(c) or _blank()
    _para(doc, f"I, {name}, hereby furnish details for ESIC registration and declare the following "
               "family members for the purpose of availing medical benefits:")
    _spacer(doc)
    _kv_table(doc, [
        ("Name:",       name),
        ("Aadhaar No.:", c.get("aadhaar_number") or _blank()),
        ("Mobile No.:", c.get("phone") or _blank()),
        ("Dispensary:", _blank()),
    ])
    _spacer(doc)
    _heading(doc, "Family Members", level=2, center=False)
    _grid_table(doc,
        ["Sr", "Name", "Relationship", "Date of Birth", "Aadhaar No."],
        [
            ["1", _blank(15), _blank(12), _blank(10), _blank(14)],
            ["2", _blank(15), _blank(12), _blank(10), _blank(14)],
            ["3", _blank(15), _blank(12), _blank(10), _blank(14)],
        ])
    _spacer(doc)
    _kv_table(doc, [("Employee Signature:", ""), ("Date:", "")])


def _s9_notice(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "घोषणा पत्र", hindi=True)
    _para(doc, "(Declaration Regarding Notice Period)", center=True, size=_SMALL_SZ, italic=True)
    _spacer(doc)
    name = _full_name(c) or _blank()
    _bilingual_para(doc,
        f"I, {name}, hereby acknowledge and agree to the Notice Period Policy of the Company:",
        f"मैं, {name}, कंपनी की नोटिस अवधि नीति को स्वीकार करता/करती हूँ:")
    _spacer(doc)
    _grid_table(doc,
        ["Sr No", "Grade", "Notice Period"],
        [
            ["1", "Trainee",                "15 Days"],
            ["2", "Probation",              "30 Days"],
            ["3", "Up to Sr. Officer",      "60 Days"],
            ["4", "Asst. Manager & Above",  "90 Days"],
        ])
    _spacer(doc)
    _bilingual_para(doc,
        "In lieu of notice, I agree to pay the Company salary for the notice period if I leave without serving the full notice.",
        "नोटिस अवधि के बदले, अगर मैं पूरी नोटिस अवधि दिए बिना कंपनी छोड़ता हूँ तो मैं नोटिस अवधि का वेतन कंपनी को देने के लिए सहमत हूँ।")
    _spacer(doc)
    _kv_table(doc, [
        ("Name:", name),
        ("Designation:", c.get("position") or _blank()),
        ("Signature:", ""),
        ("Date:", ""),
    ])


def _s10_assets(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "ASSET RESPONSIBILITY POLICY")
    _spacer(doc)
    name = _full_name(c) or _blank()
    _para(doc, f"I, {name}, acknowledge receipt of the following assets from the Company and agree to "
               "return them in the same condition at the time of my resignation or separation.")
    _spacer(doc)
    _grid_table(doc,
        ["Sr", "Asset", "Serial No / Details", "Condition", "Date Issued", "Signature"],
        [
            ["1", "Mobile Phone",   _blank(10), "Good", "", ""],
            ["2", "Laptop / Tab",   _blank(10), "Good", "", ""],
            ["3", "SIM Card",       _blank(10), "Good", "", ""],
            ["4", "ID Card",        _blank(10), "Good", "", ""],
            ["5", "Office Bag",     _blank(10), "Good", "", ""],
        ])
    _spacer(doc)
    _kv_table(doc, [
        ("Employee Name:", name),
        ("Employee Signature:", ""),
        ("Date:", ""),
        ("Issued By:", ""),
    ])


def _s11_nda(doc: Document, c: dict, company: dict):
    _page_break(doc)
    co_name = company.get("name", "Radhya Micro Finance Private Limited")
    _heading(doc, "NON-DISCLOSURE AGREEMENT (NDA)")
    _spacer(doc)
    name = _full_name(c) or _blank()
    doj  = _fmt_date(c.get("expected_joining_date"))
    _para(doc, f"This Non-Disclosure Agreement is entered into on {doj} between:")
    _spacer(doc)
    _para(doc, f"1. {co_name} ('Company')")
    _para(doc, f"2. {name} ('Employee')")
    _spacer(doc)
    clauses = [
        ("1. Confidential Information:",
         "All business strategies, client data, financial data, operational procedures, product details, and any other information not available in the public domain."),
        ("2. Obligations:",
         "The Employee agrees to keep all Confidential Information strictly confidential and not to disclose, copy, reproduce, or transmit it to any third party."),
        ("3. Duration:",
         "This NDA remains in effect during employment and for 2 (two) years after separation."),
        ("4. Consequences of Breach:",
         "Breach of this NDA will result in immediate disciplinary action including termination and may invite legal proceedings."),
        ("5. Return of Property:",
         "On separation, the Employee agrees to return all Company property and delete all confidential information from personal devices."),
    ]
    for title, body in clauses:
        _para(doc, title, bold=True)
        _para(doc, body)
    _spacer(doc)
    _sig_table(doc, ["Employee Name", "Employee Signature", "Date", "HR Signature"])


def _s12_asset_form(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "ASSET DECLARATION FORM")
    _spacer(doc)
    name = _full_name(c) or _blank()
    _para(doc, f"I, {name}, declare that the following personal assets belong to me:")
    _spacer(doc)
    _grid_table(doc,
        ["Sr", "Asset Description", "Make / Model", "Serial / IMEI", "Owned By"],
        [
            ["1", "Mobile Phone",    _blank(12), _blank(12), "Self"],
            ["2", "Laptop / Tab",    _blank(12), _blank(12), "Self"],
            ["3", "Two-Wheeler",     _blank(12), _blank(12), "Self"],
            ["4", "Other",           _blank(12), _blank(12), "Self"],
        ])
    _spacer(doc)
    _kv_table(doc, [("Employee Signature:", ""), ("Date:", "")])


def _s13_consent_docs(doc: Document, c: dict):
    _page_break(doc)
    _heading(doc, "CONSENT FORM FOR SUBMISSION OF ORIGINAL EDUCATIONAL DOCUMENTS")
    _spacer(doc)
    name = _full_name(c) or _blank()
    fhn  = c.get("father_or_husband_name") or _blank()
    doj  = _fmt_date(c.get("expected_joining_date"))
    _para(doc, f"I, {name}, son/daughter of {fhn}, hereby give my free consent to submit my original "
               "10th and 12th mark sheets to Radhya Micro Finance Pvt. Ltd. at the time of my joining.")
    _spacer(doc)
    _para(doc, "I understand and agree that:", bold=True)
    clauses = [
        "I am voluntarily submitting my original educational documents (10th & 12th mark sheets) for official verification.",
        "The Company will retain my above-mentioned original documents for a period of ONE (1) MONTH from my date of joining.",
        "After one month of joining, I will collect my original documents from the Company.",
        "The Company shall keep my documents in safe custody during the retention period.",
        "I confirm I am submitting these documents without any pressure or coercion.",
    ]
    for cl in clauses:
        _para(doc, f"• {cl}")
    _para(doc, "Documents Submitted: ORIGINAL 10TH & 12TH MARK SHEETS", bold=True)
    _spacer(doc)
    _kv_table(doc, [
        ("Date of Joining:",                doj),
        ("Date of Submission:",             ""),
        ("Submitted By (Candidate Name):",  name),
        ("Signature of Candidate:",         ""),
    ])
    _spacer(doc)
    _heading(doc, "Document Receiving Details (After 1 Month)", level=2, center=False)
    _kv_table(doc, [
        ("Date of Receiving Documents:",               ""),
        ("Received By (Candidate Name):",              ""),
        ("Signature of Receiver:",                     ""),
        ("Documents Given By (Employee Name):",        ""),
        ("Documents Given By (Employee Signature):",   ""),
    ])
    _spacer(doc)
    _heading(doc, "For Office Use Only", level=2, center=False)
    _kv_table(doc, [
        ("Authorized Signatory Name:", ""),
        ("Designation:",               ""),
        ("Signature & Stamp:",         ""),
        ("Date:",                      ""),
    ])


def _s14_posh_bgv(doc: Document, c: dict):
    _page_break(doc)
    name        = _full_name(c) or _blank()
    emp_id      = c.get("employee_id") or _blank()
    designation = c.get("position") or _blank()
    addr        = _addr_line(c) or _blank()

    # --- POSH ---
    _heading(doc, "POSH DECLARATION")
    _para(doc, "(Prevention of Sexual Harassment at Workplace)", center=True, size=_SMALL_SZ, italic=True)
    _spacer(doc)
    _para(doc, f"I, {name}, hereby declare that:")
    for cl in [
        "I have read and understood the Company's POSH (Prevention of Sexual Harassment) Policy.",
        "I am aware of the members of the Internal Complaints Committee and the process to report complaints.",
        "I agree to comply with the policy and maintain a respectful and safe workplace.",
        "I understand that any form of sexual harassment is strictly prohibited.",
        "I will report any incident of harassment that I experience or witness.",
        "I acknowledge that maintaining dignity, respect and safety at the workplace is my responsibility.",
    ]:
        _para(doc, f"• {cl}")
    _spacer(doc)
    _kv_table(doc, [
        ("Employee Name:", name),
        ("Employee ID:",   emp_id),
        ("Designation:",   designation),
        ("Signature:",     ""),
        ("Date:",          ""),
    ])
    _spacer(doc, 2)

    # --- BGV ---
    _heading(doc, "BACKGROUND VERIFICATION CONSENT FORM")
    _spacer(doc)
    _para(doc, f"I, {name}, hereby give my consent to Radhya Micro Finance Pvt. Ltd. to conduct background "
               "verification as part of my employment.")
    _para(doc, "The verification process may include:", bold=True)
    for chk in [
        "Residence / House Verification",
        "Identity & Address Verification",
        "Previous Employment Verification",
        "Police Verification (if required)",
        "Bank Statement or Financial Details (if required)",
    ]:
        _para(doc, f"• {chk}")
    _spacer(doc)
    _para(doc, "I further declare that all information provided by me is true and I authorise the Company to verify my details.", bold=True)
    _spacer(doc)
    _kv_table(doc, [
        ("Employee Name:",    name),
        ("Employee ID:",      emp_id),
        ("Designation:",      designation),
        ("Current Address:",  addr),
        ("Signature:",        ""),
        ("Date:",             ""),
    ])


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_joining_kit_docx(
    candidate: dict,
    company: dict | None = None,
    has_aadhaar_doc: bool = False,
    has_pan_doc: bool = False,
) -> bytes:
    """Build and return the Joining Kit as a .docx byte stream."""
    company = company or {}
    doc = Document()

    # Page margins (A4 — narrow margins so tables don't overflow)
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width  = Mm(210)
        section.top_margin    = Mm(18)
        section.bottom_margin = Mm(18)
        section.left_margin   = Mm(20)
        section.right_margin  = Mm(20)

    # Default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = _ENG_FONT
    style.font.size = Pt(_BODY_SZ)
    pf = style.paragraph_format
    pf.space_after  = Pt(3)
    pf.space_before = Pt(0)

    _s1_checklist(doc, candidate, company, has_aadhaar_doc, has_pan_doc)
    _s2_employee_info(doc, candidate)
    _s3_undertaking(doc, candidate)
    _s4_insurance(doc, candidate)
    _s5_gratuity(doc, candidate)
    _s6_pf(doc, candidate)
    _s7_form11(doc, candidate)
    _s8_esi(doc, candidate)
    _s9_notice(doc, candidate)
    _s10_assets(doc, candidate)
    _s11_nda(doc, candidate, company)
    _s12_asset_form(doc, candidate)
    _s13_consent_docs(doc, candidate)
    _s14_posh_bgv(doc, candidate)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
